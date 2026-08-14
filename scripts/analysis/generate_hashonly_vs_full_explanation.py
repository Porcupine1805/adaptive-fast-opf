#!/usr/bin/env python3
"""Generate the mechanistic Hash-only vs Full-FOM diagnostic report."""

from __future__ import annotations

import csv
import math
from collections import defaultdict
from datetime import date
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
RESULT_ROOT = ROOT / "results" / "experiments" / "results_full_20260811"
DIAG_DIR = RESULT_ROOT / "benchmark" / "diagnostic_mechanism_20260812"
PROBE_FILE = RESULT_ROOT / "benchmark" / "electricity_ablation_scale_probe" / "electricity_ablation_4scenario_long.csv"
EQUIV_FILE = RESULT_ROOT / "comparisons" / "full_vs_hash_only_sha256_equivalence.csv"
FIG_DIR = RESULT_ROOT / "figures" / "diagnostic_hashonly_vs_full"
REPORT_DIR = RESULT_ROOT / "reports"

CONFIG_FILES = {
    "Hash-only": DIAG_DIR / "FOMDiagnosticHashOnly_run_01.csv",
    "Hash+WSB": DIAG_DIR / "FOMDiagnosticHashWSB_run_01.csv",
    "Hash+Sparse": DIAG_DIR / "FOMDiagnosticHashSparse_run_01.csv",
    "Full-FOM": DIAG_DIR / "FOMDiagnosticFull_run_01.csv",
}

NUMERIC_FIELDS = {
    "minsup": float,
    "Time_s": float,
    "MaxMem_MB": float,
    "PairChecks": int,
    "Candidates": int,
    "Fusions": int,
    "SupportOps": int,
    "FreqPatterns": int,
    "WSBChecks": int,
    "WSBPrunes": int,
    "SparseObjects": int,
    "SparseWordsAllocated": int,
    "SparseWordsScanned": int,
    "ScalarPositionComparisons": int,
}

DATASET_LABELS = {
    "ELEC_01clients_concat.txt": "1 client",
    "ELEC_05clients_concat.txt": "5 clients",
    "ELEC_10clients_concat.txt": "10 clients",
}


def read_rows(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    for row in rows:
        for field, converter in NUMERIC_FIELDS.items():
            if field in row and row[field] != "":
                row[field] = converter(row[field])
    return rows


def load_diagnostics() -> dict[str, list[dict]]:
    data = {}
    for config, path in CONFIG_FILES.items():
        if not path.exists():
            raise FileNotFoundError(path)
        data[config] = read_rows(path)
    return data


def validate_same_semantic_work(data: dict[str, list[dict]]) -> int:
    fields = ("PairChecks", "Candidates", "Fusions", "SupportOps", "FreqPatterns")
    indexed = {}
    for config, rows in data.items():
        indexed[config] = {(row["Dataset"], row["minsup"]): row for row in rows}
    keys = set(indexed["Hash-only"])
    for config in indexed:
        if set(indexed[config]) != keys:
            raise ValueError(f"Configuration keys differ for {config}")
    matched = 0
    for key in sorted(keys):
        reference = indexed["Hash-only"][key]
        for config in indexed:
            current = indexed[config][key]
            if any(current[field] != reference[field] for field in fields):
                raise ValueError(f"Semantic work mismatch at {key}: {config}")
        matched += 1
    return matched


def summarize(data: dict[str, list[dict]]) -> list[dict]:
    summary = []
    for config, rows in data.items():
        grouped = defaultdict(list)
        for row in rows:
            grouped[row["Dataset"]].append(row)
        for dataset, values in grouped.items():
            summary.append(
                {
                    "Dataset": dataset,
                    "Config": config,
                    "MeanTime_s": sum(v["Time_s"] for v in values) / len(values),
                    "MeanMem_MB": sum(v["MaxMem_MB"] for v in values) / len(values),
                    "TotalWSBChecks": sum(v["WSBChecks"] for v in values),
                    "TotalWSBPrunes": sum(v["WSBPrunes"] for v in values),
                    "MeanPairChecks": sum(v["PairChecks"] for v in values) / len(values),
                    "MeanSparseWordsAllocated": sum(v["SparseWordsAllocated"] for v in values) / len(values),
                    "MeanSparseWordsScanned": sum(v["SparseWordsScanned"] for v in values) / len(values),
                    "MeanScalarPositionComparisons": sum(v["ScalarPositionComparisons"] for v in values) / len(values),
                    "Rows": len(values),
                }
            )
    order = {name: index for index, name in enumerate(DATASET_LABELS)}
    config_order = {name: index for index, name in enumerate(CONFIG_FILES)}
    summary.sort(key=lambda row: (order[row["Dataset"]], config_order[row["Config"]]))
    return summary


def write_csv(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def create_figures(summary: list[dict]) -> tuple[Path, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    datasets = list(DATASET_LABELS)
    x = list(range(len(datasets)))
    colors = {
        "Hash-only": "#167D6D",
        "Hash+WSB": "#4472C4",
        "Hash+Sparse": "#D9822B",
        "Full-FOM": "#B23A48",
    }
    lookup = {(row["Dataset"], row["Config"]): row for row in summary}

    runtime_path = FIG_DIR / "fig_diagnostic_runtime_by_scale.png"
    fig, ax = plt.subplots(figsize=(7.2, 4.3))
    for config in CONFIG_FILES:
        y = [lookup[(dataset, config)]["MeanTime_s"] for dataset in datasets]
        ax.plot(x, y, marker="o", linewidth=2, label=config, color=colors[config])
    ax.set_xticks(x, [DATASET_LABELS[d] for d in datasets])
    ax.set_yscale("log")
    ax.set_ylabel("Mean runtime across six minSup values (s, log scale)")
    ax.set_xlabel("ElectricityLoadDiagrams concatenation size")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=2)
    fig.tight_layout()
    fig.savefig(runtime_path, dpi=220)
    plt.close(fig)

    work_path = FIG_DIR / "fig_diagnostic_key_case.png"
    key_rows = {
        config: next(
            row for row in data_rows
            if row["Dataset"] == "ELEC_10clients_concat.txt" and row["minsup"] == 2.0
        )
        for config, data_rows in DIAGNOSTIC_DATA.items()
    }
    configs = list(CONFIG_FILES)
    fig, axes = plt.subplots(1, 2, figsize=(8.0, 4.1))
    axes[0].bar(configs, [key_rows[c]["Time_s"] for c in configs], color=[colors[c] for c in configs])
    axes[0].set_ylabel("Runtime (s)")
    axes[0].set_title("Runtime")
    axes[0].tick_params(axis="x", rotation=25)
    axes[0].grid(axis="y", alpha=0.2)
    allocated_gib = [key_rows[c]["SparseWordsAllocated"] * 8 / (1024 ** 3) for c in configs]
    axes[1].bar(configs, allocated_gib, color=[colors[c] for c in configs])
    axes[1].set_ylabel("Cumulative long[] payload allocated (GiB)")
    axes[1].set_title("Bitmap allocation traffic")
    axes[1].tick_params(axis="x", rotation=25)
    axes[1].grid(axis="y", alpha=0.2)
    fig.suptitle("10 clients, minSup = 2")
    fig.tight_layout()
    fig.savefig(work_path, dpi=220)
    plt.close(fig)
    return runtime_path, work_path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "D9EAF0")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    return table


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_code_reference(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(65, 65, 65)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def find_probe_values() -> dict[tuple[str, str, float], dict]:
    rows = read_rows(PROBE_FILE)
    return {(row["Config"], row["Dataset"], row["minsup"]): row for row in rows}


def count_equivalence() -> tuple[int, int]:
    with EQUIV_FILE.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    return sum(row["status"] == "SHA256_MATCH" for row in rows), len(rows)


def generate_report(data: dict[str, list[dict]], summary: list[dict], runtime_fig: Path, work_fig: Path) -> Path:
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Giải thích cơ chế: Vì sao Hash-only nhanh hơn Full-FOM")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Đối chiếu benchmark, bộ đếm thực thi và implementation hiện tại").italic = True
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(f"Ngày lập báo cáo: {date.today().strftime('%d/%m/%Y')}")

    document.add_heading("Kết luận điều hành", level=1)
    document.add_paragraph(
        "Hash-only nhanh hơn Full-FOM không phải vì tối ưu hash làm thay đổi kết quả khai phá, mà vì trong "
        "implementation đang được benchmark, hash join đã loại phần lớn cặp không tương thích, trong khi hai lớp "
        "bổ sung của Full-FOM không tạo thêm mức giảm công việc đủ để bù chi phí của chúng. Cụ thể, WSB hiện cắt 0 "
        "cặp do vị trí và công thức kiểm tra làm điều kiện cắt không thể xảy ra; còn cấu trúc mang tên SparseBitmap "
        "thực tế cấp phát một dải word liên tục từ minBlock đến maxBlock, sao chép/shift các dải này và quét cả các "
        "word bằng 0. Trên chuỗi dài, chi phí cấp phát, quét bộ nhớ và GC vượt xa lợi ích của phép AND theo word."
    )
    document.add_paragraph(
        "Vì vậy, số liệu hiện tại chỉ chứng minh mạnh đóng góp của hash-indexed join. Chúng chưa chứng minh rằng kiến "
        "trúc Full-FOM được mô tả trong bản thảo (sparse non-zero-word, fused zero-allocation kernel và residual "
        "pruning) đã được hiện thực và đánh giá đúng. Claim trong bài cần được thu hẹp hoặc code phải được viết lại "
        "đúng thiết kế rồi benchmark lại."
    )

    document.add_heading("1. Thiết kế kiểm tra", level=1)
    document.add_paragraph(
        "Bốn cấu hình được chạy trong các JVM riêng với cùng input ElectricityLoadDiagrams, cùng sáu ngưỡng "
        "minSup = {2, 4, 6, 8, 10, 12}, cùng heap -Xms2g -Xmx16g. Mỗi cấu hình có 18 ca: 1, 5 và 10 client "
        "nhân sáu ngưỡng. Đây là chạy chẩn đoán một lần để giải thích cơ chế; số runtime không thay thế benchmark "
        "nhiều lần dùng cho kiểm định thống kê."
    )
    add_table(
        document,
        ["Cấu hình", "Hash", "SparseBitmap", "WSB", "Đường fusion"],
        [
            ["Hash-only", "Bật", "Tắt", "Tắt", "Scalar two-pointer"],
            ["Hash+WSB", "Bật", "Tắt", "Bật", "Scalar two-pointer"],
            ["Hash+Sparse", "Bật", "Bật", "Tắt", "Bitmap shift/intersection"],
            ["Full-FOM", "Bật", "Bật", "Bật", "Bitmap shift/intersection"],
        ],
    )
    document.add_paragraph(
        "Bộ đếm được thêm vào đường thực thi nhưng không thay đổi điều kiện sinh pattern hoặc support. Runtime "
        "không instrumentation từ đợt scale probe được dùng như bằng chứng hiệu năng chính; lượt instrumented dùng "
        "để định vị nguồn chi phí."
    )

    document.add_heading("2. Bằng chứng định lượng", level=1)
    lookup = {(row["Dataset"], row["Config"]): row for row in summary}
    runtime_rows = []
    for dataset in DATASET_LABELS:
        hash_time = lookup[(dataset, "Hash-only")]["MeanTime_s"]
        runtime_rows.append(
            [
                DATASET_LABELS[dataset],
                f"{hash_time:.3f}",
                f"{lookup[(dataset, 'Hash+WSB')]['MeanTime_s']:.3f}",
                f"{lookup[(dataset, 'Hash+Sparse')]['MeanTime_s']:.3f}",
                f"{lookup[(dataset, 'Full-FOM')]['MeanTime_s']:.3f}",
                f"{lookup[(dataset, 'Full-FOM')]['MeanTime_s'] / hash_time:.2f}x",
            ]
        )
    add_table(
        document,
        ["Quy mô", "Hash-only (s)", "Hash+WSB (s)", "Hash+Sparse (s)", "Full-FOM (s)", "Full/Hash"],
        runtime_rows,
    )
    caption = document.add_paragraph("Hình 1. Runtime trung bình của lượt chẩn đoán (trục log).")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_picture(str(runtime_fig), width=Inches(6.2))

    probe = find_probe_values()
    p_hash = probe[("HashOnly", "ELEC_10clients_concat.txt", 2.0)]
    p_full = probe[("Full", "ELEC_10clients_concat.txt", 2.0)]
    document.add_paragraph(
        f"Ở lượt không instrumentation đã chạy trước đó, ca 10 clients/minSup=2 cho kết quả Hash-only "
        f"{p_hash['Time_s']:.3f} s và Full-FOM {p_full['Time_s']:.3f} s, tức Full chậm hơn "
        f"{p_full['Time_s'] / p_hash['Time_s']:.2f}x. Điều này xác nhận chênh lệch không phải do riêng overhead "
        "của bộ đếm chẩn đoán."
    )

    key = {
        config: next(
            row for row in rows
            if row["Dataset"] == "ELEC_10clients_concat.txt" and row["minsup"] == 2.0
        )
        for config, rows in data.items()
    }
    key_rows = []
    for config in CONFIG_FILES:
        row = key[config]
        key_rows.append(
            [
                config,
                f"{row['Time_s']:.3f}",
                f"{row['MaxMem_MB']:.0f}",
                f"{row['PairChecks']:,}",
                f"{row['WSBPrunes']:,}",
                f"{row['SparseWordsAllocated'] * 8 / (1024 ** 3):.2f}",
                f"{row['SparseWordsScanned'] / 1e9:.3f}B",
                f"{row['ScalarPositionComparisons'] / 1e6:.1f}M",
            ]
        )
    add_table(
        document,
        ["Cấu hình", "Time (s)", "Heap MB*", "PairChecks", "WSB prune", "long[] GiB**", "Word scan", "Scalar cmp"],
        key_rows,
    )
    document.add_paragraph(
        "* Heap MB là used-heap được lấy mẫu, không phải RSS hay allocation profiling. ** long[] GiB là tổng payload "
        "8 byte/word được cấp phát lũy kế; đây là lưu lượng cấp phát, không phải bộ nhớ đồng thời còn sống."
    )
    document.add_picture(str(work_fig), width=Inches(6.4))
    caption = document.add_paragraph("Hình 2. Ca chẩn đoán 10 clients, minSup=2.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    matched_cases = validate_same_semantic_work(data)
    sha_matches, sha_total = count_equivalence()
    document.add_paragraph(
        f"Cả bốn cấu hình có PairChecks, Candidates, Fusions, SupportOps và FreqPatterns giống nhau trong "
        f"{matched_cases}/{matched_cases} ca điện lực. Ngoài ra, kiểm tra canonical đã có trước đó cho Full-FOM và "
        f"Hash-only đạt SHA-256 raw exact match {sha_matches}/{sha_total} file. Như vậy khác biệt runtime không đến "
        "từ việc Hash-only sinh ít output hơn."
    )

    document.add_heading("3. Vì sao WSB không giúp trong code hiện tại", level=1)
    document.add_paragraph(
        "Trước khi vào WSB, code đã bỏ qua p nếu p.prefixSupport < minSup và bỏ qua q nếu q.suffixSupport < "
        "minSup. Khi tới WSB, ta luôn có p.prefixSupport >= minSup và q.suffixSupport >= minSup. Vì k = 1/n > 0 "
        "nên exp(k) > 1, do đó p.prefixSupport*exp(k) >= minSup. Cả hai đối số của min đều >= minSup, vì vậy:"
    )
    formula = document.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.add_run("upperBound = min(p.prefixSupport × exp(k), q.suffixSupport) ≥ minSup").bold = True
    document.add_paragraph(
        "Điều kiện upperBound < minSup không thể đúng trong số học hữu hạn thông thường. Bộ đếm xác nhận tổng "
        "1,874,790 lần kiểm tra WSB trên 18 ca của Hash+WSB nhưng WSBPrunes = 0; Full-FOM cũng bằng 0. Vì thế "
        "Hash+WSB thực hiện cùng semantic work với Hash-only, còn khác biệt thời gian nhỏ giữa hai lượt là nhiễu "
        "JIT/GC/OS của phép đo một lần, không phải speedup do pruning."
    )
    add_code_reference(document, "FOMAblationFlags.java:231-250 (matchingQs, support guards và WSB bound)")
    document.add_paragraph(
        "Điểm sâu hơn là WSB trong bài được mô tả như residual-support bound trên phần occurrence chưa xử lý. Code "
        "hiện tại chỉ dùng hai scalar support tổng tại đầu cặp; nó không cập nhật cận theo cursor/word còn lại trong "
        "kernel. Đây không phải cùng một cơ chế với Algorithm 3 được mô tả trong manuscript."
    )

    document.add_heading("4. Vì sao SparseBitmap làm Full-FOM chậm", level=1)
    add_bullet(
        document,
        "Representation không sparse theo non-zero word: create() cấp phát long[maxBlock-minBlock+1]. Chỉ cần occurrence đầu và cuối cách xa nhau thì toàn bộ khoảng giữa được giữ, kể cả phần lớn word bằng 0.",
    )
    add_bullet(
        document,
        "Mỗi generation giữ bitmap trong PatternCandidate. Với nhiều pattern có span dài, live heap tăng theo số pattern nhân với độ dài span/64, không theo số non-zero word như claim trong bài.",
    )
    add_bullet(
        document,
        "Đường nóng tạo object/mảng trung gian: getBitmap().shiftForwardOneStep() cấp phát long[] mới cho p; q.getBitmap().copy() cấp phát long[] mới cho q; mỗi fuse còn tạo hai ArrayList cho rPositions/hPositions.",
    )
    add_bullet(
        document,
        "fuseVector() duyệt tuần tự mọi block trong overlapMin..overlapMax. Khi occurrence thưa nhưng span rộng, nó quét nhiều zero-region. Ca 10 clients/minSup=2 quét 7.330 tỷ word để thu được cùng 18.898 triệu SupportOps.",
    )
    add_bullet(
        document,
        "Hash-only dùng hai con trỏ trên int[] occurrence đã sắp thứ tự. Nó chỉ thực hiện khoảng 112.2 triệu bước so sánh ở ca trên, truy cập tuyến tính, không materialize bitmap và không phát sinh hàng chục GiB allocation traffic.",
    )
    add_code_reference(document, "FOMAblationFlags.java:70-106 (range bitmap, copy và shift allocation)")
    add_code_reference(document, "FOMAblationFlags.java:264-271, 279-310 (chọn vector/scalar và vòng quét block)")

    document.add_heading("5. Vì sao ba tối ưu không cộng dồn tuyến tính", level=1)
    document.add_paragraph(
        "Các option nằm ở các công đoạn khác nhau nhưng không độc lập về chi phí. Hash join thay đổi phân bố công "
        "việc đi vào các công đoạn sau: sau khi cặp không tương thích đã bị loại, tập cặp còn lại đủ nhỏ để scalar "
        "two-pointer xử lý hiệu quả. Lúc này bitmap chỉ có lợi nếu mật độ occurrence đủ cao và implementation bỏ qua "
        "được zero-region với chi phí cấp phát thấp. Code hiện tại không đạt hai điều kiện đó. Tương tự, pruning chỉ "
        "có lợi nếu cận đủ chặt và được cập nhật khi phần residual giảm; cận hiện tại không cắt gì. Do đó tổng thời "
        "gian có dạng T_full = T_hash + overhead_bitmap + overhead_WSB - work_saved. Trong phép đo hiện tại, "
        "work_saved của WSB bằng 0 và lợi ích bitmap âm, nên T_full > T_hash-only."
    )
    document.add_paragraph(
        "Điều này không chứng minh rằng sparse word kernel hoặc residual pruning đúng nghĩa luôn vô ích. Nó chứng "
        "minh rằng phiên bản hiện tại chưa hiện thực các cơ chế đó theo cách đủ để tạo lợi ích, và dữ liệu lớn hơn "
        "không tự động sửa được vấn đề: range bitmap còn tăng chi phí theo span khi n tăng."
    )

    document.add_heading("6. Khoảng cách giữa manuscript và code", level=1)
    add_table(
        document,
        ["Mô tả trong bài", "Code đang chạy", "Hệ quả"],
        [
            ["Chỉ lưu non-zero words và word indices", "Một long[] liên tục minBlock..maxBlock", "Không có độ phức tạp theo số non-zero word; tốn RAM khi span rộng"],
            ["Fused shift-intersection không tạo object trung gian", "shiftForwardOneStep(), copy() và ArrayList cấp phát trong hot path", "Allocation traffic và GC lớn"],
            ["Residual-support bound trên phần chưa xử lý", "min của hai support tổng, kiểm tra sau guards >= minSup", "Điều kiện prune bất khả thi; 0/1,874,790 lần cắt"],
            ["Full-FOM giảm heap", "Full dùng heap cao hơn Hash-only trên electricity probe", "Không thể dùng kết quả hiện tại để xác nhận memory claim của Full"],
            ["Cumulative ablation cải thiện tuần tự", "Hash+Sparse gần Full và chậm hơn Hash-only rõ rệt", "Table/claim ablation phải chạy lại sau khi code khớp thuật toán"],
        ],
    )

    document.add_heading("7. Nội dung đề xuất sửa trong bài báo", level=1)
    document.add_paragraph(
        "Đoạn dưới đây có thể dùng để thay thế phần diễn giải ablation hiện tại cho đến khi Full-FOM được hiện thực "
        "lại và benchmark đầy đủ."
    )
    proposed = (
        "Contrary to our initial expectation, enabling all three optimization switches did not yield the lowest "
        "runtime in the current implementation. Across the ElectricityLoadDiagrams scale probe, the hash-only "
        "configuration consistently outperformed the full configuration while producing the same candidate, "
        "fusion, support-operation, and frequent-pattern counts. Previously generated canonical outputs also "
        "showed exact SHA-256 agreement between hash-only and full configurations in all 84 evaluated files.\n\n"
        "Mechanistic profiling attributes this result to implementation overhead rather than an inherent negative "
        "interaction among the intended algorithms. First, the implemented weighted-support bound is evaluated "
        "only after both residual support variables have already been constrained to be at least minSup. Since "
        "exp(k) is greater than one, the resulting bound cannot fall below minSup; accordingly, no pair was pruned "
        "in 1,874,790 instrumented checks. Second, the current bitmap structure stores every machine word between "
        "the first and last occupied blocks rather than storing only non-zero words. It also materializes copied "
        "and shifted arrays in the fusion path. For the 10-client, minSup=2 case, this path scanned 7.330 billion "
        "bitmap words and generated approximately 98.2 GiB of cumulative long-array payload allocation, whereas "
        "the hash-only scalar merge required 112.2 million ordered-position comparisons.\n\n"
        "These findings establish the hash-indexed join as the dominant validated optimization in the evaluated "
        "artifact. They do not invalidate a genuinely sparse, allocation-aware shift-intersection kernel or a "
        "cursor-updated residual-support bound. However, claims concerning those mechanisms and the combined "
        "Full-FOM configuration must be treated as design hypotheses until the implementation stores only non-zero "
        "words, removes per-fusion bitmap materialization, applies a decreasing residual bound inside the processing "
        "loop, and is re-evaluated with repeated, order-balanced benchmarks and allocation/GC profiling."
    )
    for block in proposed.split("\n\n"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(block)
        run.italic = True

    document.add_heading("8. Hành động cần thiết trước khi giữ claim Full-FOM", level=1)
    add_bullet(document, "Thay range bitmap bằng cặp sparse (wordIndex[], wordMask[]) chỉ chứa word khác 0, hoặc dùng container bitmap đã được kiểm chứng.")
    add_bullet(document, "Fuse shift và intersection bằng cursor trên hai sparse-word arrays; không tạo shifted/copy long[] cho từng p-q pair.")
    add_bullet(document, "Đưa residual bound vào vòng xử lý, cập nhật theo phần endpoint/weight chưa tiêu thụ; thêm test bắt buộc WSBPrunes > 0 trên dữ liệu được thiết kế để prune.")
    add_bullet(document, "Chạy factorial 2^3 đầy đủ, nhiều lần, đổi thứ tự cấu hình theo round, warm-up rõ ràng và báo median/IQR hoặc mean/CI.")
    add_bullet(document, "Đo allocation bytes, GC pause/count và live heap bằng JFR/GC logs; không suy ra zero-allocation chỉ từ source hoặc sampled used heap.")
    add_bullet(document, "Lặp canonical pattern-support SHA-256 trên electricity và mọi dataset sau mỗi thay đổi implementation.")

    document.add_heading("9. Artefact và khả năng tái lập", level=1)
    artefacts = [
        "src/benchmark/java/FOMAblationFlags.java",
        "results/experiments/results_full_20260811/benchmark/diagnostic_mechanism_20260812/",
        "results/experiments/results_full_20260811/benchmark/electricity_ablation_scale_probe/",
        "results/experiments/results_full_20260811/comparisons/full_vs_hash_only_sha256_equivalence.csv",
        "scripts/benchmark/run_benchmark.ps1",
        "scripts/analysis/generate_hashonly_vs_full_explanation.py",
    ]
    for item in artefacts:
        add_code_reference(document, item)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    output = REPORT_DIR / "hash_only_vs_full_fom_mechanistic_explanation.docx"
    document.save(output)
    return output


if __name__ == "__main__":
    DIAGNOSTIC_DATA = load_diagnostics()
    same_work_cases = validate_same_semantic_work(DIAGNOSTIC_DATA)
    SUMMARY = summarize(DIAGNOSTIC_DATA)
    write_csv(DIAG_DIR / "diagnostic_summary_by_dataset.csv", SUMMARY)
    key_case_rows = []
    for config, rows in DIAGNOSTIC_DATA.items():
        for row in rows:
            if row["Dataset"] == "ELEC_10clients_concat.txt" and row["minsup"] == 2.0:
                key_case_rows.append({"Config": config, **row})
    write_csv(DIAG_DIR / "diagnostic_key_case_10clients_minsup2.csv", key_case_rows)
    runtime_figure, work_figure = create_figures(SUMMARY)
    report_path = generate_report(DIAGNOSTIC_DATA, SUMMARY, runtime_figure, work_figure)
    print(f"Validated identical semantic metrics: {same_work_cases} cases")
    print(f"Report: {report_path}")
