#!/usr/bin/env python3
"""Create a manuscript revision package that remains faithful to the current code."""

from __future__ import annotations

import csv
import math
import statistics
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "results" / "experiments" / "results_full_20260811"
OPF_CSV = RESULTS / "benchmark" / "opf" / "OPF_Miner_Original_summary_avg.csv"
HASH_CSV = RESULTS / "benchmark" / "ablation_clean" / "FOMAblationHashOnly_summary_avg.csv"
FULL_CSV = RESULTS / "benchmark" / "ablation_clean" / "FOMAblationFull_summary_avg.csv"
ELECTRICITY_CSV = RESULTS / "benchmark" / "electricity_ablation_scale_probe" / "electricity_ablation_4scenario_long.csv"
DIAGNOSTIC_CSV = RESULTS / "benchmark" / "diagnostic_mechanism_20260812" / "diagnostic_summary_by_dataset.csv"
OPF_FULL_EQUIV = RESULTS / "comparisons" / "sha256_equivalence.csv"
FULL_HASH_EQUIV = RESULTS / "comparisons" / "full_vs_hash_only_sha256_equivalence.csv"
OUTPUT = RESULTS / "reports" / "manuscript_revision_keep_code_scientific.docx"


def rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def indexed(path: Path) -> dict[tuple[str, float], dict[str, str]]:
    return {(r["Dataset"], float(r["minsup"])): r for r in rows(path)}


def geo_mean(values: list[float]) -> float:
    return math.exp(sum(math.log(v) for v in values) / len(values))


OPF = indexed(OPF_CSV)
HASH = indexed(HASH_CSV)
FULL = indexed(FULL_CSV)
KEYS = sorted(set(OPF) & set(HASH))

SPEEDUPS = [float(OPF[k]["Time_s"]) / float(HASH[k]["Time_s"]) for k in KEYS]
MEMORY_RATIOS = [float(OPF[k]["MaxMem_MB"]) / float(HASH[k]["MaxMem_MB"]) for k in KEYS]
FULL_HASH_RATIOS = [float(FULL[k]["Time_s"]) / float(HASH[k]["Time_s"]) for k in KEYS]
FUSION_REDUCTIONS = [1.0 - float(HASH[k]["Fusions"]) / float(OPF[k]["Fusions"]) for k in KEYS]
SUPPORT_REDUCTIONS = [1.0 - float(HASH[k]["SupportOps"]) / float(OPF[k]["SupportOps"]) for k in KEYS]

TIME_WINS = sum(v > 1.0 for v in SPEEDUPS)
MEMORY_WINS = sum(v > 1.0 for v in MEMORY_RATIOS)
FULL_SLOWER_CASES = sum(v > 1.0 for v in FULL_HASH_RATIOS)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], data: list[list[str]], font_size: float = 8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade(cell, "D9EAF0")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)
    for values in data:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return table


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_replacement(doc: Document, heading: str, text: str) -> None:
    doc.add_heading(heading, level=2)
    for block in text.strip().split("\n\n"):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block.strip())
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        shade_paragraph(paragraph, "F3F6F7")


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    for style_name, size in (("Normal", 10.5), ("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)


def exactness_counts() -> tuple[Counter, int, int]:
    opf_full = rows(OPF_FULL_EQUIV)
    full_hash = rows(FULL_HASH_EQUIV)
    return Counter(r["status"] for r in opf_full), len(opf_full), sum(r["status"] == "SHA256_MATCH" for r in full_hash)


def electricity_summary() -> dict[tuple[str, str], dict[str, float]]:
    grouped: dict[tuple[str, str], list[dict[str, str]]] = {}
    for row in rows(ELECTRICITY_CSV):
        grouped.setdefault((row["Dataset"], row["Config"]), []).append(row)
    result = {}
    for key, values in grouped.items():
        result[key] = {
            "time": statistics.mean(float(v["Time_s"]) for v in values),
            "memory": statistics.mean(float(v["MaxMem_MB"]) for v in values),
            "pairchecks": statistics.mean(float(v["PairChecks"]) for v in values),
        }
    return result


def add_metric_summary(doc: Document) -> None:
    doc.add_heading("3. Evidence that remains defensible", level=1)
    add_table(
        doc,
        ["Evidence", "Observed result", "Permissible claim"],
        [
            ["Runtime: OPF vs Hash-only", f"{TIME_WINS}/84 wins; geometric mean {geo_mean(SPEEDUPS):.2f}x; median {statistics.median(SPEEDUPS):.2f}x; maximum {max(SPEEDUPS):.2f}x", "Hash indexing is the dominant validated acceleration in this artifact"],
            ["Count equivalence", "84/84 configurations have matching frequent-pattern counts", "No cardinality loss was observed"],
            ["Canonical equivalence", "OPF vs Full: 77 raw SHA, 6 normalized SHA, 1 tolerance match, 0 failures; Full vs Hash-only: 84/84 raw SHA", "Hash-only is equivalent to OPF under the documented comparison protocol"],
            ["Fusion work", f"Counter reduction range {min(FUSION_REDUCTIONS)*100:.1f}%–{max(FUSION_REDUCTIONS)*100:.1f}%", "The index avoids most broad pair-processing work"],
            ["Support work", f"Counter reduction range {min(SUPPORT_REDUCTIONS)*100:.1f}%–{max(SUPPORT_REDUCTIONS)*100:.1f}%", "Candidate localization reduces downstream support work in the measured implementation"],
            ["Sampled used heap", f"Hash-only lower in {MEMORY_WINS}/84 cases; geometric mean OPF/Hash ratio {geo_mean(MEMORY_RATIOS):.2f}x", "Report only as a sampled JVM proxy, not peak-live memory or allocation proof"],
            ["Full vs Hash-only", f"Full slower in {FULL_SLOWER_CASES}/84 original cases; geometric mean Full/Hash ratio {geo_mean(FULL_HASH_RATIOS):.2f}x", "Optimization composition is non-monotonic in the current implementation"],
            ["Electricity scale probe", "Hash-only fastest in 18/18 configurations", "The bitmap path does not become favorable merely because input length grows"],
            ["WSB diagnostic", "1,874,790 checks and 0 prunes", "No runtime benefit can be attributed to WSB in its current placement"],
        ],
        7.8,
    )


def add_minsup4_table(doc: Document) -> None:
    data = []
    for dataset in [f"DB{i}.txt" for i in range(1, 9)]:
        key = (dataset, 4.0)
        opf, hsh = OPF[key], HASH[key]
        speedup = float(opf["Time_s"]) / float(hsh["Time_s"])
        data.append([
            dataset.removesuffix(".txt"),
            f"{float(opf['Time_s']):.6f}",
            f"{float(hsh['Time_s']):.6f}",
            f"{speedup:.2f}x",
            f"{int(float(opf['Fusions'])):,}",
            f"{int(float(hsh['Fusions'])):,}",
            hsh["FreqPatterns"],
        ])
    add_table(
        doc,
        ["Dataset", "OPF (s)", "FastOPF-HJ (s)", "Speedup", "OPF fusion counter", "HJ fusion counter", "Patterns"],
        data,
        7.7,
    )


def add_db5_table(doc: Document) -> None:
    data = []
    for minsup in (2.0, 4.0, 6.0, 8.0, 10.0, 12.0):
        key = ("DB5.txt", minsup)
        opf, hsh = OPF[key], HASH[key]
        speedup = float(opf["Time_s"]) / float(hsh["Time_s"])
        data.append([
            f"{minsup:g}", f"{float(opf['Time_s']):.6f}", f"{float(hsh['Time_s']):.6f}",
            f"{speedup:.2f}x", hsh["FreqPatterns"],
        ])
    add_table(doc, ["minSup", "OPF (s)", "FastOPF-HJ (s)", "Speedup", "Patterns"], data)


def add_electricity_table(doc: Document) -> None:
    elec = electricity_summary()
    data = []
    labels = {
        "ELEC_01clients_concat.txt": "1 client",
        "ELEC_05clients_concat.txt": "5 clients",
        "ELEC_10clients_concat.txt": "10 clients",
    }
    for dataset, label in labels.items():
        h = elec[(dataset, "HashOnly")]
        f = elec[(dataset, "Full")]
        data.append([
            label, f"{h['time']:.3f}", f"{f['time']:.3f}", f"{f['time']/h['time']:.2f}x",
            f"{h['memory']:.1f}", f"{f['memory']:.1f}", f"{f['memory']/h['memory']:.1f}x",
        ])
    add_table(
        doc,
        ["Scale", "Hash-only (s)", "Full (s)", "Full/Hash time", "Hash MB", "Full MB", "Full/Hash memory"],
        data,
    )


doc = Document()
configure(doc)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Gói nội dung sửa manuscript khi giữ nguyên code hiện tại")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Scientific repositioning of FastOPF around the validated hash-indexed contribution").italic = True
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Generated on {date.today().strftime('%d/%m/%Y')} from the current benchmark artifacts")

doc.add_heading("1. Recommended scientific position", level=1)
doc.add_paragraph(
    "Có thể giữ nguyên code và vẫn duy trì một đóng góp đáng kể, nhưng đối tượng được xem là thuật toán đề xuất phải "
    "chuyển từ Full-FOM sang cấu hình hash-indexed scalar fusion. Tên khuyến nghị là FastOPF-HJ. Sparse range "
    "bitmap và WSB tiếp tục tồn tại trong code như các experimental switches, nhưng không được mô tả là các tối ưu "
    "đã được xác nhận. Full-FOM trở thành một composite prototype trong ablation study, không phải cấu hình đại diện."
)
doc.add_paragraph(
    "Điểm mới trung tâm vẫn rõ ràng: candidate fusion của OPF có cấu trúc equijoin giữa normalized suffix và prefix. "
    "Lập chỉ mục theo prefix signature chuyển broad pair scan thành output-sensitive compatible-pair enumeration, "
    "giữ nguyên semantics và tạo mức giảm công việc lớn. Đây là đóng góp thuật toán và systems độc lập với bitmap "
    "hoặc pruning."
)
add_table(
    doc,
    ["Recommended title", "Use"],
    [
        ["FastOPF-HJ: Exact Output-Sensitive Hash-Indexed Fusion for Forgetting-Aware Order-Preserving Pattern Mining", "Khuyến nghị mạnh nhất; tên phản ánh đúng implementation được bảo vệ"],
        ["FastOPF-Miner: Exact Hash-Indexed Candidate Fusion for Forgetting-Aware Order-Preserving Pattern Mining", "Giữ thương hiệu FastOPF nhưng bỏ allocation-aware/sparse claims"],
        ["When More Optimizations Are Slower: Exact Hash-Indexed OPF Mining and an Empirical Study of Optimization Composition", "Phù hợp venue systems/experimental; nhấn mạnh negative result"],
    ],
)
doc.add_heading("Artifact-to-manuscript naming", level=2)
add_table(
    doc,
    ["Manuscript name", "Artifact entry point", "Role"],
    [
        ["OPF-Miner", "OPF_Miner_Original.java", "Reference baseline"],
        ["FastOPF-HJ", "FOMAblationFlags.java with -Dmode=hash_only", "Proposed and headline configuration"],
        ["Hash+WSB", "FOMAblationFlags.java with -Dmode=hash_wsb", "Experimental interaction test"],
        ["Hash+RangeBitmap", "FOMAblationFlags.java with -Dmode=hash_sparse", "Experimental representation test; do not call genuinely sparse"],
        ["Full composite", "FOMAblationFlags.java with -Dmode=full", "Ablation configuration, not the proposed fastest method"],
        ["Legacy FOM executable", "FOM.java", "Retained artifact; do not mix its timing rows with the FastOPF-HJ headline table"],
    ],
    7.8,
)

doc.add_heading("2. Claim audit: keep, rewrite, or remove", level=1)
add_table(
    doc,
    ["Current manuscript claim", "Decision", "Scientifically safe replacement"],
    [
        ["Output-sensitive hash-indexed join", "KEEP", "Primary contribution; retain completeness and expected output-sensitive complexity analysis"],
        ["Sparse structure stores only non-zero words", "REMOVE", "Current structure stores a contiguous minBlock-to-maxBlock word range"],
        ["Fused kernel eliminates transient allocation", "REMOVE", "Current path copies/materializes arrays and lists; no zero-allocation claim"],
        ["Residual pruning aggressively skips pairs", "REMOVE AS BENEFIT", "Describe the implemented guard only as an evaluated switch; diagnostic observed zero prunes"],
        ["Full configuration is the proposed fastest method", "REWRITE", "FastOPF-HJ/hash-only is the validated proposed configuration"],
        ["8.8x geometric mean and 271.3x peak", "REPLACE", f"Current fair summaries: {geo_mean(SPEEDUPS):.2f}x geometric mean, {statistics.median(SPEEDUPS):.2f}x median, {max(SPEEDUPS):.2f}x peak for FastOPF-HJ"],
        ["Memory reduction up to 9.2x validates zero allocation", "REWRITE", f"Sampled used heap is lower in {MEMORY_WINS}/84 cases; do not infer allocation behavior"],
        ["Cumulative ablation improves monotonically", "REMOVE", "Factorial results show non-monotonic interactions and representation-dependent overhead"],
        ["Strict output equivalence", "KEEP WITH PROTOCOL", "Report raw SHA, normalized SHA and tolerance categories separately"],
        ["Hardware-conscious architecture", "REWRITE", "Use implementation-conscious/output-sensitive; avoid cache/SIMD claims without counters"],
    ],
    7.7,
)

add_metric_summary(doc)
doc.add_heading("Headline table at minSup = 4", level=2)
add_minsup4_table(doc)
doc.add_heading("DB5 sensitivity using current fair summaries", level=2)
add_db5_table(doc)
doc.add_heading("Electricity scale probe: why Full is not the representative method", level=2)
add_electricity_table(doc)
doc.add_paragraph(
    "The ElectricityLoadDiagrams values are one-run scale-probe measurements and should be used to explain the "
    "mechanism and motivate limitations, not as confidence-interval benchmark results."
)

doc.add_heading("4. Replacement-ready English manuscript text", level=1)

abstract = f"""
Order-preserving pattern mining with an exponential forgetting mechanism identifies recurrent temporal shapes while assigning greater importance to recent observations. The reference OPF-Miner workflow, however, repeatedly considers broad sets of candidate pairs even though a fusion is possible only when the normalized suffix of one pattern equals the normalized prefix of another. We present FastOPF-HJ, an exact reformulation that treats candidate fusion as an output-sensitive equijoin. At each pattern length, frequent patterns are indexed by normalized prefix signatures, and each suffix query retrieves only structurally compatible partners. Hash collisions cannot affect correctness because candidate keys are verified by structural equality, while occurrence fusion and forgetting-weighted support retain the reference semantics.

Across 84 dataset-threshold configurations, FastOPF-HJ reproduced the reference frequent-pattern count in every case and was consistent with the documented canonical-output comparison protocol. It reduced the implementation's fusion-operation counter by {min(FUSION_REDUCTIONS)*100:.1f}% to {max(FUSION_REDUCTIONS)*100:.1f}% and its support-operation counter by {min(SUPPORT_REDUCTIONS)*100:.1f}% to {max(SUPPORT_REDUCTIONS)*100:.1f}%. Using the available ten-run summaries, FastOPF-HJ was faster than OPF-Miner in {TIME_WINS} of 84 configurations, with a geometric-mean speedup of {geo_mean(SPEEDUPS):.2f}x, a median speedup of {statistics.median(SPEEDUPS):.2f}x, and a maximum observed speedup of {max(SPEEDUPS):.2f}x. A factorial implementation study further shows that additional bitmap and support-bound switches do not improve performance monotonically: in the current artifact, the hash-only configuration is faster than the full composite in {FULL_SLOWER_CASES} of 84 original configurations and all 18 ElectricityLoadDiagrams scale-probe configurations. Mechanistic counters attribute this outcome to contiguous-range bitmap traffic and an upper-bound guard that pruned no pairs. These results establish hash-indexed compatible-pair enumeration as the dominant validated optimization and demonstrate the importance of reporting optimization interactions rather than assuming their benefits are additive.
"""
add_replacement(doc, "4.1 Revised abstract", abstract)

keywords = "Order-preserving pattern mining; forgetting-aware support; output-sensitive algorithms; hash join; exact pattern mining; ablation study; reproducible benchmarking"
add_replacement(doc, "4.2 Revised keywords", keywords)

contributions = """
The main contributions of this work are as follows.

(1) Exact hash-indexed compatible-pair enumeration. We formulate prefix-suffix candidate fusion as an equijoin over normalized overlap signatures. A prefix hash index retrieves only structurally compatible partners, while explicit array equality preserves correctness under hash collisions.

(2) Output-sensitive complexity characterization. For a generation containing f_m frequent patterns of length m and J_m compatible ordered pairs, index construction and lookup require expected O(m f_m) work, and enumeration requires O(J_m) work. The J_m term is unavoidable for an exact algorithm that explicitly processes every compatible pair.

(3) Semantics-preserving integration with forgetting-aware occurrence fusion. The indexed schedule changes only how compatible parents are located. Rank fusion, endpoint alignment, exponential weights, and the minimum weighted-support criterion remain unchanged, enabling direct equivalence validation against OPF-Miner.

(4) Reproducible exactness and work-reduction evaluation. We report frequent-pattern counts, canonical pattern-support comparison categories, operation counters, and raw per-configuration measurements instead of relying only on aggregate speedups.

(5) Factorial analysis of optimization composition. We show empirically that separately motivated optimizations need not compose monotonically. In the evaluated Java artifact, contiguous-range bitmap processing introduces allocation and scan costs after hash indexing has already reduced the candidate set, while the implemented support-bound guard does not prune any pair. This negative result provides practical guidance for representation selection and for the evaluation of composite mining systems.
"""
add_replacement(doc, "4.3 Revised contribution list", contributions)

design = """
FastOPF-HJ preserves the mining semantics of OPF-Miner and changes the schedule used to locate fusion-compatible parents. For each pattern length, the algorithm first computes the normalized prefix and suffix signatures of every frequent pattern. Patterns are inserted into a hash map keyed by their normalized prefix. For each parent p, querying the map with suffix(p) returns the list of q patterns satisfying suffix(p) = prefix(q). The retrieved pairs then pass through the same group-validity test, occurrence alignment, boundary comparison, forgetting-weight accumulation, and minimum-support test used by the reference formulation.

This separation is important. Hash indexing reduces structural search but does not change the number of truly compatible pairs, the occurrence endpoints generated from a compatible pair, or the support assigned to a child. The implementation uses immutable integer-array keys with cached hash codes and full structural equality. Consequently, collisions may increase lookup cost but cannot introduce an invalid pair or remove a valid pair.

Occurrence lists in the FastOPF-HJ configuration are sorted integer endpoints. To fuse p and q, the endpoint array of p is shifted logically by one and intersected with the endpoint array of q using a two-pointer merge. Every aligned endpoint is decoded according to the OPF fusion rule and contributes its precomputed exponential weight exactly once. This scalar merge is the primary evaluated occurrence representation; no sparse-word, SIMD, or zero-allocation property is required by the proposed method.
"""
add_replacement(doc, "4.4 Revised Section IV-A: Design overview", design)

theorem = """
Theorem (Completeness and exactness of indexed pair enumeration). Let F_m be the set of frequent patterns of length m, and let H map every normalized prefix signature to all patterns in F_m having that prefix. For every ordered pair (p,q) in F_m x F_m, q is returned by H[suffix(p)] if and only if suffix(p) = prefix(q). Therefore, the indexed schedule enumerates exactly the structurally joinable ordered pairs considered by OPF-Miner after its overlap-equality test. Since the subsequent group test, endpoint fusion, and weighted-support calculation are unchanged, FastOPF-HJ produces the same child candidates, occurrence endpoints, and support values as the reference algorithm, subject to the same strict-tie policy.

Complexity. Constructing normalized keys for F_m requires O(m|F_m|) time. Under expected constant-time hash operations, index insertion and lookup require O(|F_m|), and enumerating the returned lists requires O(J_m), where J_m is the number of compatible ordered pairs. The expected structural-join cost is therefore O(m|F_m| + J_m), with O(m|F_m|) key storage. This bound is output-sensitive: no exact explicit-fusion algorithm can avoid O(J_m) work when all J_m pairs must be processed.
"""
add_replacement(doc, "4.5 Revised theorem and complexity statement", theorem)

methodology = """
We evaluate OPF-Miner and FastOPF-HJ on the same input files and minimum-support thresholds. The current result archive contains ten complete runs per configuration for the original DB1-DB8 benchmark suite. Each CSV records runtime, sampled JVM used heap, candidate/fusion counters, support-operation counters, and frequent-pattern counts. Because sub-100-ms measurements are sensitive to JIT compilation, garbage collection, operating-system scheduling, and fixed execution order, we report the number of wins together with geometric means, medians, per-configuration values, and standard deviations rather than interpreting every small timing difference as an algorithmic effect.

Correctness is evaluated at two levels. First, frequent-pattern cardinalities are compared for every dataset-threshold configuration. Second, canonical sorted pattern-support outputs are compared by raw SHA-256, normalized SHA-256 after documented numeric normalization, and an explicitly reported tolerance category. These categories are not conflated: a tolerance match is evidence of numeric agreement under the selected threshold, not a byte-identical file.

The ElectricityLoadDiagrams experiment is reported separately as a scale probe. It uses concatenations of 1, 5, and 10 client series and six minimum-support thresholds. Because only one complete run is currently available per configuration, this experiment is used for mechanism diagnosis and external-validity discussion rather than formal timing inference. Hardware, operating-system, JVM build, heap flags, and CPU-governor metadata must be inserted from the actual execution environment before submission.
"""
add_replacement(doc, "4.6 Revised experimental methodology", methodology)

results_text = f"""
FastOPF-HJ matched the OPF-Miner frequent-pattern count in all 84 dataset-threshold configurations. The canonical comparison between OPF-Miner and the full implementation produced 77 raw SHA-256 matches, six normalized SHA-256 matches, one tolerance match, and no failures. The independently generated full and hash-only canonical files matched by raw SHA-256 in all 84 cases. Together, these checks support equivalence of FastOPF-HJ under the documented comparison protocol while preserving the distinction between byte-level and tolerance-level agreement.

The primary work reduction is attributable to candidate localization. Relative to the OPF counters, the hash-indexed configuration reduced counted fusion operations by {min(FUSION_REDUCTIONS)*100:.1f}% to {max(FUSION_REDUCTIONS)*100:.1f}% and support operations by {min(SUPPORT_REDUCTIONS)*100:.1f}% to {max(SUPPORT_REDUCTIONS)*100:.1f}% across the evaluated configurations. These reductions should be interpreted as implementation counters, not hardware instructions or cache-miss measurements.

FastOPF-HJ was faster in {TIME_WINS} of 84 configurations. The geometric-mean OPF/FastOPF-HJ runtime ratio was {geo_mean(SPEEDUPS):.2f}x, the median was {statistics.median(SPEEDUPS):.2f}x, and the largest observed ratio was {max(SPEEDUPS):.2f}x on DB5 at minSup=2. The method was slower in {84-TIME_WINS} configurations, with the minimum observed ratio of {min(SPEEDUPS):.2f}x. This qualification matters because the strongest gains occur where broad candidate enumeration is expensive, whereas very small configurations may be dominated by index construction, JVM, and measurement overhead.

Sampled JVM used heap was lower for FastOPF-HJ in {MEMORY_WINS} of 84 configurations, with a geometric-mean OPF/FastOPF-HJ ratio of {geo_mean(MEMORY_RATIOS):.2f}x. This measure is a coarse occupancy proxy. It does not establish peak live memory, allocation rate, garbage-collection time, native memory, or RSS, and no zero-allocation claim is made.
"""
add_replacement(doc, "4.7 Revised exactness, work, runtime, and memory results", results_text)

ablation = f"""
The factorial ablation contradicts the assumption that enabling every optimization switch must produce the fastest implementation. On the original 84 configurations, hash-only was faster than the full composite in {FULL_SLOWER_CASES} cases, and the geometric-mean Full/Hash-only runtime ratio was {geo_mean(FULL_HASH_RATIOS):.2f}x. On the ElectricityLoadDiagrams scale probe, hash-only was fastest in all 18 configurations and the gap widened with concatenation length.

Mechanistic instrumentation explains this result. The current bitmap class allocates a contiguous long-array spanning minBlock through maxBlock; it does not store only non-zero words. Its fusion path creates copied or shifted arrays and scans every overlapping word, including zero regions. At 10 clients and minSup=2, the full path scanned 7.330 billion bitmap words and generated approximately 98.2 GiB of cumulative long-array payload allocation, while hash-only used 112.2 million ordered-position comparisons. In the same diagnostic suite, the weighted support-bound switch executed 1,874,790 checks and pruned zero pairs. The guard is reached only after both support variables have been constrained to be at least minSup; because exp(k)>1, min(pSupport*exp(k),qSupport) cannot fall below minSup at that point.

These observations do not show that sparse bitmaps or residual bounds are intrinsically ineffective. They show that benefits depend on representation density, allocation policy, bound tightness, and placement in the execution schedule. The current evidence therefore validates hash indexing but treats genuinely sparse non-zero-word storage and cursor-updated residual pruning as future implementation directions rather than completed contributions.
"""
add_replacement(doc, "4.8 Revised ablation and mechanism analysis", ablation)

discussion = """
The contribution of FastOPF-HJ is deliberately narrower than a complete redesign of occurrence processing. It establishes that the prefix-suffix compatibility condition can be exposed as an exact output-sensitive hash join and demonstrates that this scheduling change accounts for the dominant work reduction in the current artifact. The exponential forgetting model, rank semantics, fusion cases, and weighted-support definition remain inherited from OPF-Miner.

The experiments also provide a broader systems lesson: optimization layers at different pipeline stages are not performance-independent. Candidate indexing changes the workload presented to occurrence fusion. Once incompatible pairs have been removed, a sequential merge over compact integer arrays can be preferable to materializing and scanning range bitmaps. Similarly, an upper bound has value only if it decreases with unprocessed work and is evaluated before its condition becomes redundant. Reporting these interactions is a scientific result because it identifies the workload and implementation conditions under which common optimization intuitions fail.

Several limitations remain. The worst-case output is still combinatorial; hash indexing cannot avoid enumerating compatible pairs. The model uses strict ranks and rejects ties. Most original datasets are financial, and the ElectricityLoadDiagrams scale probe currently lacks repeated trials. Runtime summaries are affected by JVM and system noise, especially below 100 ms. Used-heap sampling is not an allocation profile. Finally, the current range bitmap and support-bound switches should not be interpreted as implementations of a non-zero-word sparse kernel or a dynamically decreasing residual-support algorithm.
"""
add_replacement(doc, "4.9 Revised discussion and limitations", discussion)

conclusion = f"""
This study presented FastOPF-HJ, an exact output-sensitive hash-indexed schedule for forgetting-aware order-preserving pattern mining. By indexing normalized prefix signatures and querying them with normalized suffix signatures, the method avoids broad verification of structurally incompatible candidate pairs while preserving OPF-Miner's rank, occurrence, and weighted-support semantics. Across 84 evaluated dataset-threshold configurations, FastOPF-HJ matched the reference frequent-pattern counts, reduced measured fusion and support work, and achieved a geometric-mean speedup of {geo_mean(SPEEDUPS):.2f}x, with wins in {TIME_WINS} configurations.

The factorial study further demonstrated that optimization benefits are not additive by default. In the current Java artifact, contiguous-range bitmap processing and a non-pruning support guard outweighed their intended benefits after hash indexing had already reduced the candidate workload. We therefore limit the validated contribution to exact hash-indexed enumeration and report the additional switches as implementation findings. Future work will investigate adaptive occurrence representations, genuinely sparse non-zero-word fusion, and cursor-updated residual bounds under repeated allocation-aware benchmarks.
"""
add_replacement(doc, "4.10 Revised conclusion", conclusion)

doc.add_heading("5. Required manuscript surgery by section", level=1)
add_table(
    doc,
    ["Location in FastOPF-V1.2", "Action", "Replacement source in this document"],
    [
        ["Title", "Replace allocation-aware title", "Section 1 title options"],
        ["Abstract", "Replace completely; remove 8.8x/271.3x/9.2x and sparse/zero-allocation claims", "Section 4.1"],
        ["Keywords", "Remove JVM memory optimization; add hash join and ablation", "Section 4.2"],
        ["Introduction contributions 2-5", "Delete sparse, fused kernel, WSB, allocation-contract contributions", "Section 4.3"],
        ["Section IV-A", "Reframe method around indexed schedule", "Section 4.4"],
        ["Section IV-B", "Retain hash join; ensure notation renders correctly", "Sections 4.4-4.5"],
        ["Sections IV-C to IV-E", "Move current bitmap/WSB to implementation variants; remove theorems 2-3, Proposition 1, and pruning corollary as claims about artifact", "Section 4.8"],
        ["Section IV-F", "Replace end-to-end proof with indexed-schedule theorem", "Section 4.5"],
        ["Tables II-III", "Redesign for FastOPF-HJ scalar path; remove zero-allocation and sparse-block rows", "Sections 4.4-4.5"],
        ["Methodology", "Remove unsupported multi-fork, warm-up, CPU-frequency claims unless logs prove them; insert real metadata", "Section 4.6"],
        ["Table V", "Use OPF, Hash-only/FastOPF-HJ, Hash+WSB, Hash+Sparse, Full; label latter three experimental", "Section 4.8"],
        ["Tables VI-VII", "Retain/recalculate; separate cardinality from canonical protocol and label counters exactly", "Sections 4.7 and 3"],
        ["Tables VIII-IX", "Replace with current fair OPF vs Hash-only numbers", "Headline and DB5 tables in Section 3"],
        ["Table X", "Recalculate OPF vs Hash-only; call sampled used heap, not footprint reduction proof", "Section 4.7"],
        ["Table XI", "Do not retain old scaling claim unless regenerated for FastOPF-HJ under the same protocol", "Mark pending"],
        ["Table XII", "Delete cumulative monotonic ablation; replace with factorial and Electricity analysis", "Sections 3 and 4.8"],
        ["Discussion", "Replace contribution boundary and add implementation gap", "Section 4.9"],
        ["Conclusion", "Replace completely", "Section 4.10"],
    ],
    7.5,
)

doc.add_heading("6. Language that must not remain", level=1)
for phrase in [
    "stores only non-zero machine words",
    "zero allocation post-initialization",
    "entirely eliminates transient object allocation",
    "aggressive residual-support pruning",
    "SIMD/vectorization (unless hardware instructions are measured)",
    "full configuration systematically/always outperforms hash-only",
    "memory reduction validates the allocation contract",
    "successive integration systematically diminishes execution time",
    "8.8x geometric mean, 271.3x peak, and 9.2x memory reduction from the old manuscript tables",
]:
    add_bullet(doc, phrase)

doc.add_heading("7. Submission checklist", level=1)
for item in [
    "Insert actual machine, OS, JVM build, heap flags, and run dates; do not retain placeholder Ubuntu/i7 metadata unless it is the execution machine.",
    "Regenerate Tables VIII-X from OPF_Miner_Original_summary_avg.csv and FOMAblationHashOnly_summary_avg.csv.",
    "Preserve all 84 raw per-run CSVs and report that summaries contain ten runs.",
    "Keep SHA comparison categories separate: raw, normalized, tolerance, failure.",
    "State that the Electricity scale probe has one run per configuration and use it only as diagnostic evidence.",
    "Rename the representative implementation/configuration consistently in code appendix, tables, figures, and captions.",
    "Have a second author cross-check every numeric claim against the generated tables before submission.",
]:
    add_number(doc, item)

doc.add_heading("8. Source artifacts used", level=1)
for artifact in [OPF_CSV, HASH_CSV, FULL_CSV, ELECTRICITY_CSV, DIAGNOSTIC_CSV, OPF_FULL_EQUIV, FULL_HASH_EQUIV]:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(artifact.relative_to(ROOT)).replace("\\", "/"))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(70, 70, 70)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
print(f"Paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
